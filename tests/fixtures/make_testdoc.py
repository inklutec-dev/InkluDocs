#!/usr/bin/env python3
"""Erzeugt ein FIKTIVES Word-Testdokument fuer das InkluDocs-Word-Werkzeug.

Abgedeckte Faelle (jeder Fall ist im Dokument selbst als Text benannt):
  A  Inline-Bild im Fliesstext mit Bildunterschrift (Formatvorlage Caption)
  B  Bild mit VORHANDENEM Alt-Text (descr) und Titel
  C  Bild, das bereits als DEKORATIV markiert ist (adec:decorative)
  D  Bild in einer Tabellenzelle
  E  Bild in der Kopfzeile (Logo)
  F  Frei positioniertes Bild (wp:anchor statt wp:inline)
  G  JPEG statt PNG
  A2 Dasselbe Bild wie A ein zweites Mal (gleicher Medienpart -> Dedup)
Alle Inhalte sind erfunden (Testdaten, keine echten Personen/Firmen).
"""
from pathlib import Path
from PIL import Image, ImageDraw
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import copy

OUT = Path(__file__).parent
IMG = OUT / "img"; IMG.mkdir(exist_ok=True)

def bild(name, farbe, text, fmt="PNG", size=(640, 400)):
    im = Image.new("RGB", size, farbe)
    d = ImageDraw.Draw(im)
    d.rectangle([20, 20, size[0]-20, size[1]-20], outline="white", width=6)
    d.text((40, 40), text, fill="white")
    d.text((40, 80), "Testbild (fiktiv)", fill="white")
    p = IMG / f"{name}.{fmt.lower()}"
    im.save(p, fmt)
    return p

A = bild("bild_a", (40, 90, 160), "Bild A: Balkendiagramm-Attrappe")
B = bild("bild_b", (160, 60, 40), "Bild B: Logo-Attrappe mit Alt-Text")
C = bild("bild_c", (90, 90, 90), "Bild C: Zierlinie (dekorativ)", size=(640, 60))
D = bild("bild_d", (40, 140, 90), "Bild D: Foto-Attrappe in Tabelle")
E = bild("bild_e", (20, 20, 60), "Bild E: Kopfzeilen-Logo", size=(300, 100))
F = bild("bild_f", (150, 120, 30), "Bild F: frei positioniert")
G = bild("bild_g", (120, 40, 120), "Bild G: JPEG-Foto-Attrappe", fmt="JPEG")

doc = Document()
doc.core_properties.title = "Testdokument Word-Werkzeug (fiktiv)"
doc.core_properties.author = "InkluDocs Testgenerator"

# Kopfzeile mit Logo (Fall E)
hdr = doc.sections[0].header
hp = hdr.paragraphs[0]
hp.add_run().add_picture(str(E), width=Cm(3))
hp.add_run("  Musterfirma Beispiel GmbH (fiktiv)")

doc.add_heading("Testdokument für das InkluDocs-Word-Werkzeug", level=0)
doc.add_paragraph("Dieses Dokument ist vollständig erfunden und dient nur dem Test der "
                  "Alt-Text-Erkennung. Es enthält Bilder an verschiedenen Stellen.")

doc.add_heading("1 Einleitung", level=1)
doc.add_paragraph("Die Umsatzentwicklung der fiktiven Musterfirma wird in Abbildung 1 "
                  "als Balkendiagramm dargestellt. Die Werte sind ausgedacht.")
doc.add_paragraph().add_run().add_picture(str(A), width=Cm(12))          # Fall A
cap = doc.add_paragraph("Abbildung 1: Umsatz der Musterfirma 2023 bis 2026 (fiktive Werte)", style="Caption")
doc.add_paragraph("Wie Abbildung 1 zeigt, steigt der Umsatz in jedem Jahr. Dieser Absatz "
                  "steht nach der Bildunterschrift und liefert weiteren Kontext.")

doc.add_heading("1.1 Vorhandener Alt-Text", level=2)
doc.add_paragraph("Das folgende Bild trägt bereits einen Alternativtext und einen Titel, "
                  "die das Werkzeug als Vorlage erkennen soll.")
run_b = doc.add_paragraph().add_run(); run_b.add_picture(str(B), width=Cm(6))   # Fall B
docPr_b = run_b._r.xpath(".//wp:docPr")[0]
docPr_b.set("descr", "Logo der Musterfirma: roter Kreis auf weißem Grund (alter Alt-Text, fiktiv)")
docPr_b.set("title", "Musterfirma-Logo")

doc.add_heading("1.2 Dekoratives Bild", level=2)
doc.add_paragraph("Die Zierlinie darunter ist in Word bereits als dekorativ markiert.")
run_c = doc.add_paragraph().add_run(); run_c.add_picture(str(C), width=Cm(12))  # Fall C
docPr_c = run_c._r.xpath(".//wp:docPr")[0]
ext = parse_xml(
    '<a:extLst xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
    '<a:ext uri="{C183D7F6-B498-43B3-948B-1728B52AA6E4}">'
    '<adec:decorative xmlns:adec="http://schemas.microsoft.com/office/drawing/2017/decorative" val="1"/>'
    '</a:ext></a:extLst>')
docPr_c.append(ext)

doc.add_heading("2 Tabelle mit Bild", level=1)
doc.add_paragraph("In der Tabelle steht ein Bild in der zweiten Zelle (Fall D).")
tbl = doc.add_table(rows=2, cols=2); tbl.style = "Table Grid"
tbl.cell(0, 0).text = "Produkt (fiktiv)"; tbl.cell(0, 1).text = "Foto"
tbl.cell(1, 0).text = "Beispielgerät Modell X"
tbl.cell(1, 1).paragraphs[0].add_run().add_picture(str(D), width=Cm(4))        # Fall D

doc.add_heading("3 Frei positioniertes Bild", level=1)
doc.add_paragraph("Das nächste Bild ist nicht im Textfluss verankert, sondern frei "
                  "positioniert (wp:anchor). Solche Bilder werden gern übersehen.")
run_f = doc.add_paragraph().add_run(); run_f.add_picture(str(F), width=Cm(8))   # Fall F
inline = run_f._r.xpath(".//wp:inline")[0]
drawing = inline.getparent()
anchor = parse_xml(
    '<wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
    'distT="0" distB="0" distL="114300" distR="114300" simplePos="0" relativeHeight="251658240" '
    'behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">'
    '<wp:simplePos x="0" y="0"/>'
    '<wp:positionH relativeFrom="column"><wp:posOffset>914400</wp:posOffset></wp:positionH>'
    '<wp:positionV relativeFrom="paragraph"><wp:posOffset>0</wp:posOffset></wp:positionV>'
    '</wp:anchor>')
for child in list(inline):            # extent, effectExtent, docPr, cNvGraphicFramePr, graphic
    anchor.append(child)
# wrapSquare muss vor docPr stehen (Schema-Reihenfolge): extent, effectExtent, wrap*, docPr ...
wrap = parse_xml('<wp:wrapSquare xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" wrapText="bothSides"/>')
idx = [c.tag for c in anchor].index(qn("wp:docPr"))
anchor.insert(idx, wrap)
drawing.replace(inline, anchor)

doc.add_heading("4 JPEG und Wiederholung", level=1)
doc.add_paragraph("Ein JPEG-Foto (Fall G) und danach noch einmal Bild A (Fall A2), "
                  "das denselben Medienpart nutzt.")
doc.add_paragraph().add_run().add_picture(str(G), width=Cm(8))               # Fall G
doc.add_paragraph("Abbildung 2: Fiktives Produktfoto", style="Caption")
doc.add_paragraph().add_run().add_picture(str(A), width=Cm(8))               # Fall A2
doc.add_paragraph("Abbildung 3: Dieselbe Grafik wie Abbildung 1, erneut eingefügt", style="Caption")

doc.add_heading("5 Schluss", level=1)
doc.add_paragraph("Ende des Testdokuments. Alle Angaben sind erfunden.")

out = OUT / "testdokument_inkludocs.docx"
doc.save(out)
print("geschrieben:", out, out.stat().st_size, "Bytes")

# Negativ-Faelle: Altformat und Makro-Datei (nur zum Abweisen)
(OUT / "altformat.doc").write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 512)
import shutil; shutil.copy(out, OUT / "makro_test.docm")
# Kein-Word-Zip: eine Zip-Datei ohne word/document.xml
import zipfile
with zipfile.ZipFile(OUT / "kein_word.docx", "w") as z:
    z.writestr("hallo.txt", "kein Word-Dokument")
print("Negativ-Dateien geschrieben")
